#!/usr/bin/env python3
"""
bygg-rapport.py — Generisk mall för att bygga docx-rapporter åt
sammanfatta-media-skillen.

Använd så här:
    1. Anropa bygg_rapport(data, output_path) med en strukturerad dict.
    2. Funktionen producerar en .docx med IBM Plex Sans, A4, 1-tums marginaler.
    3. Smart-quote-konvertering och zoom-attribute-fix körs automatiskt.

Beroenden:
    pip install python-docx

Validering (efter att rapporten är klar):
    python3 ~/.claude/skills/docx/scripts/office/validate.py <fil>

Om validering klagar på saknad w:percent på w:zoom: zoom-fixen körs som
post-process i denna fil, men kräver att docx-skillens unpack.py/pack.py
finns tillgängliga. Om de saknas: fixet hoppas över utan fel, och Word
kommer ändå att kunna öppna filen — det är bara schema-validatorn som
klagar.

Utvecklat 2026-05 utifrån Heiser- och Bickle-körningarna.
"""

from __future__ import annotations

import shutil
import subprocess
import sys
import tempfile
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

FONT = "IBM Plex Sans"
LEFT_QUOTE = "“"   # "
RIGHT_QUOTE = "”"  # "


# -----------------------------------------------------------------------------
# Lågnivå-hjälpare för typsnitt och sidlayout
# -----------------------------------------------------------------------------

def _set_default_font(doc: Document, font_name: str = FONT, size: int = 11) -> None:
    """Sätt default-typsnitt för hela dokumentet."""
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(size)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for k in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(k), font_name)


def _apply_font(run, font_name: str = FONT) -> None:
    """Sätt typsnitt på en specifik run."""
    run.font.name = font_name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for k in ("w:ascii", "w:hAnsi", "w:cs"):
        rFonts.set(qn(k), font_name)


def _set_a4(doc: Document) -> None:
    """A4-sidstorlek med 2,54 cm marginaler."""
    for section in doc.sections:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)


def _style_headings(doc: Document) -> None:
    """Sätt IBM Plex Sans och mörk färg på alla rubrikstilar."""
    for level in (1, 2, 3):
        style = doc.styles[f"Heading {level}"]
        style.font.name = FONT
        style.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
        rPr = style.element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        rFonts.set(qn("w:ascii"), FONT)
        rFonts.set(qn("w:hAnsi"), FONT)


# -----------------------------------------------------------------------------
# Innehållshjälpare
# -----------------------------------------------------------------------------

def add_para(doc: Document, text: str):
    """Lägg till en löpande paragraf."""
    p = doc.add_paragraph()
    r = p.add_run(text)
    _apply_font(r)
    return p


def add_labeled_para(doc: Document, label: str, body: str):
    """Paragraf med fet etikett följt av brödtext."""
    p = doc.add_paragraph()
    r1 = p.add_run(label)
    _apply_font(r1)
    r1.bold = True
    r2 = p.add_run(body)
    _apply_font(r2)
    return p


def add_bullet(doc: Document, text: str):
    """Punktlistsrad."""
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    _apply_font(r)
    return p


def add_labeled_bullet(doc: Document, label: str, body: str):
    """Punktlistsrad med fet etikett."""
    p = doc.add_paragraph(style="List Bullet")
    r1 = p.add_run(label)
    _apply_font(r1)
    r1.bold = True
    r2 = p.add_run(body)
    _apply_font(r2)
    return p


def add_heading(doc: Document, text: str, level: int = 1):
    """Wrapper kring add_heading som applicerar IBM Plex Sans."""
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        _apply_font(r)
        r.bold = True
    return h


# -----------------------------------------------------------------------------
# Post-processorer (smart quotes, zoom-fix)
# -----------------------------------------------------------------------------

def _convert_quotes_in_run(text: str, state: dict) -> str:
    """Par-baserad konvertering av raka " till smart quotes."""
    out = []
    for c in text:
        if c == '"':
            if state["open"]:
                out.append(RIGHT_QUOTE)
                state["open"] = False
            else:
                out.append(LEFT_QUOTE)
                state["open"] = True
        else:
            out.append(c)
    return "".join(out)


def apply_smart_quotes(doc: Document) -> dict:
    """Konvertera alla raka " till typografiska. Returnerar slut-state för debug."""
    state = {"open": False}
    for p in doc.paragraphs:
        for run in p.runs:
            run.text = _convert_quotes_in_run(run.text, state)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.text = _convert_quotes_in_run(run.text, state)
    return state


def fix_zoom_attribute(docx_path: Path) -> bool:
    """Fixa python-docx-buggen där <w:zoom> saknar w:percent."""
    candidate_paths = [
        Path.home() / ".claude" / "skills" / "docx" / "scripts" / "office",
        Path("/mnt/skills/public/docx/scripts/office"),
    ]
    skill_office = next((p for p in candidate_paths if (p / "unpack.py").exists()), None)
    if skill_office is None:
        print("INFO: docx-skillens unpack.py hittades inte; hoppar över zoom-fix.",
              file=sys.stderr)
        return False

    with tempfile.TemporaryDirectory(prefix="docx-zoom-fix-") as tmpdir:
        work = Path(tmpdir) / "unpacked"
        try:
            subprocess.run(
                ["python3", str(skill_office / "unpack.py"), str(docx_path), str(work)],
                check=True, capture_output=True,
            )
            settings = work / "word" / "settings.xml"
            content = settings.read_text(encoding="utf-8")
            new_content = content.replace(
                '<w:zoom w:val="bestFit"/>',
                '<w:zoom w:val="bestFit" w:percent="100"/>',
            )
            if new_content == content:
                return True
            settings.write_text(new_content, encoding="utf-8")
            subprocess.run(
                ["python3", str(skill_office / "pack.py"),
                 str(work), str(docx_path), "--original", str(docx_path)],
                check=True, capture_output=True,
            )
            return True
        except subprocess.CalledProcessError as exc:
            print(f"VARNING: zoom-fix misslyckades: {exc}", file=sys.stderr)
            return False


# -----------------------------------------------------------------------------
# Huvudfunktion
# -----------------------------------------------------------------------------

def bygg_rapport(data: dict, output_path: str | Path) -> Path:
    """Bygg en docx-rapport från strukturerad data.

    Förväntad datastruktur:
        {
            "titel": "Rapportens huvudrubrik",
            "metadata": [("Källa: ", "..."), ("URL: ", "..."), ...],
            "sektioner": [
                {
                    "rubrik": "Översikt",
                    "stycken": ["paragraf 1", "paragraf 2", ...],
                    "bullets": [("Etikett: ", "text"), ...],   # frivilligt
                    "underrubriker": [                          # frivilligt
                        {
                            "rubrik": "Tes",
                            "stycken": [...],
                            "bullets": [...],
                        },
                    ],
                },
            ],
        }
    """
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    _set_a4(doc)
    _set_default_font(doc)
    _style_headings(doc)

    # Titel
    add_heading(doc, data["titel"], level=0)

    # Metadatarader
    for label, body in data.get("metadata", []):
        add_labeled_para(doc, label, body)
    add_para(doc, "")

    # Sektioner
    for sektion in data.get("sektioner", []):
        add_heading(doc, sektion["rubrik"], level=1)
        for stycke in sektion.get("stycken", []):
            add_para(doc, stycke)
        for bullet in sektion.get("bullets", []):
            if isinstance(bullet, tuple):
                add_labeled_bullet(doc, bullet[0], bullet[1])
            else:
                add_bullet(doc, bullet)
        for under in sektion.get("underrubriker", []):
            add_heading(doc, under["rubrik"], level=2)
            for stycke in under.get("stycken", []):
                add_para(doc, stycke)
            for bullet in under.get("bullets", []):
                if isinstance(bullet, tuple):
                    add_labeled_bullet(doc, bullet[0], bullet[1])
                else:
                    add_bullet(doc, bullet)

    doc.save(str(output_path))

    # Post-processorer
    doc2 = Document(str(output_path))
    state = apply_smart_quotes(doc2)
    doc2.save(str(output_path))
    if state["open"]:
        print("VARNING: obalanserade citattecken (oddat antal \").", file=sys.stderr)

    fix_zoom_attribute(output_path)

    return output_path


# -----------------------------------------------------------------------------
# Exempel
# -----------------------------------------------------------------------------

EXEMPELDATA = {
    "titel": "Exempelrapport — minimal mall",
    "metadata": [
        ("Källa: ", "YouTube"),
        ("URL: ", "https://www.youtube.com/watch?v=EXEMPEL"),
        ("Hämtad: ", "2026-05-13"),
        ("Rapportnivå: ", "0 — ingen diagnostik"),
    ],
    "sektioner": [
        {
            "rubrik": "Sammanfattning",
            "underrubriker": [
                {
                    "rubrik": "Tes",
                    "stycken": [
                        'Videons huvudpåstående. Citat med "raka tecken" konverteras automatiskt.'
                    ],
                },
            ],
        },
        {
            "rubrik": "Källor",
            "underrubriker": [
                {
                    "rubrik": "Bibelreferenser",
                    "bullets": [
                        ("Ps 82 — ", "Exempelreferens."),
                    ],
                },
            ],
        },
    ],
}

if __name__ == "__main__":
    out = Path(sys.argv[1]) if len(sys.argv) > 1 else Path("rapport-exempel.docx")
    result = bygg_rapport(EXEMPELDATA, out)
    print(f"Skapade: {result}")
